"""
Publishing module for BlueWriter.
Handles PDF, DOCX, and EPUB generation for rough and final drafts.
"""
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from html.parser import HTMLParser

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, 
    Table, TableStyle
)
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ebooklib import epub

from models.story import Story
from models.chapter import Chapter
from database.connection import DatabaseManager


class HTMLTextExtractor(HTMLParser):
    """Extract plain text from HTML, preserving basic structure."""
    
    def __init__(self):
        super().__init__()
        self.result = []
        self.current_tag = None
    
    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        if tag in ('p', 'div', 'br'):
            pass  # Will add newline on end tag
        elif tag in ('h1', 'h2', 'h3', 'h4'):
            self.result.append('\n\n')
    
    def handle_endtag(self, tag):
        if tag in ('p', 'div', 'br'):
            self.result.append('\n')
        elif tag in ('h1', 'h2', 'h3', 'h4'):
            self.result.append('\n')
        self.current_tag = None
    
    def handle_data(self, data):
        self.result.append(data)
    
    def get_text(self) -> str:
        return ''.join(self.result).strip()


def html_to_text(html: str) -> str:
    """Convert HTML to plain text."""
    if not html:
        return ""
    parser = HTMLTextExtractor()
    parser.feed(html)
    return parser.get_text()


def html_to_paragraphs(html: str) -> List[str]:
    """Convert HTML to list of paragraph strings."""
    text = html_to_text(html)
    # Split on double newlines or single newlines
    paragraphs = re.split(r'\n\s*\n|\n', text)
    return [p.strip() for p in paragraphs if p.strip()]


class StoryPublisher:
    """Handles publishing stories to various formats."""
    
    def __init__(self, db_path: str):
        self.db_path = db_path
        self.styles = getSampleStyleSheet()
        self._setup_styles()
    
    def _setup_styles(self) -> None:
        """Set up custom paragraph styles."""
        # Title style
        self.styles.add(ParagraphStyle(
            name='BookTitle',
            parent=self.styles['Title'],
            fontSize=28,
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        
        # Chapter title style
        self.styles.add(ParagraphStyle(
            name='ChapterTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceBefore=20,
            spaceAfter=20,
            alignment=TA_CENTER
        ))
        
        # Body text style
        self.styles.add(ParagraphStyle(
            name='StoryBody',
            parent=self.styles['Normal'],
            fontSize=12,
            leading=18,
            spaceBefore=0,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            firstLineIndent=24
        ))
        
        # First paragraph (no indent)
        self.styles.add(ParagraphStyle(
            name='StoryBodyFirst',
            parent=self.styles['StoryBody'],
            firstLineIndent=0
        ))
        
        # Synopsis style
        self.styles.add(ParagraphStyle(
            name='Synopsis',
            parent=self.styles['Normal'],
            fontSize=11,
            leading=16,
            spaceBefore=10,
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Times-Italic'
        ))
    
    def publish_to_pdf(self, story_id: int, output_path: str, 
                       include_synopsis: bool = True,
                       draft_watermark: bool = False) -> str:
        """
        Publish a story to PDF format.
        
        Args:
            story_id: ID of the story to publish
            output_path: Path for the output PDF file
            include_synopsis: Whether to include chapter summaries
            draft_watermark: Whether to add "DRAFT" watermark
            
        Returns:
            Path to the created PDF file
        """
        with DatabaseManager(self.db_path) as conn:
            story = Story.get_by_id(conn, story_id)
            if not story:
                raise ValueError(f"Story with ID {story_id} not found")
            
            chapters = Chapter.get_by_story(conn, story_id)
        
        # Sort chapters by board position (left to right, top to bottom)
        chapters.sort(key=lambda c: (c.board_y, c.board_x))
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=1*inch,
            leftMargin=1*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        # Build content
        content = []
        
        # Title page
        content.append(Spacer(1, 2*inch))
        content.append(Paragraph(story.title, self.styles['BookTitle']))
        
        if story.synopsis:
            content.append(Spacer(1, 0.5*inch))
            content.append(Paragraph(story.synopsis, self.styles['Synopsis']))
        
        if draft_watermark:
            content.append(Spacer(1, 1*inch))
            content.append(Paragraph(
                f"DRAFT - Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                self.styles['Synopsis']
            ))
        
        content.append(PageBreak())
        
        # Chapters
        for i, chapter in enumerate(chapters):
            # Chapter title
            chapter_num = i + 1
            title_text = f"Chapter {chapter_num}: {chapter.title}"
            content.append(Paragraph(title_text, self.styles['ChapterTitle']))
            
            # Chapter summary (optional)
            if include_synopsis and chapter.summary:
                content.append(Paragraph(
                    f"<i>{chapter.summary}</i>", 
                    self.styles['Synopsis']
                ))
            
            # Chapter content
            paragraphs = html_to_paragraphs(chapter.content)
            for j, para in enumerate(paragraphs):
                # Use non-indented style for first paragraph
                style = self.styles['StoryBodyFirst'] if j == 0 else self.styles['StoryBody']
                content.append(Paragraph(para, style))
            
            # Page break between chapters (except last)
            if i < len(chapters) - 1:
                content.append(PageBreak())
        
        # Build the PDF
        doc.build(content)
        
        return output_path

    def publish_rough_draft(self, story_id: int, output_dir: str) -> str:
        """
        Publish a rough draft PDF.
        Includes draft watermark and timestamp.
        Does NOT lock the story.
        
        Returns:
            Path to the created PDF file
        """
        with DatabaseManager(self.db_path) as conn:
            story = Story.get_by_id(conn, story_id)
            if not story:
                raise ValueError(f"Story with ID {story_id} not found")
            
            # Update status to rough published
            story.publish_rough(conn)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_title = re.sub(r'[^\w\s-]', '', story.title).strip().replace(' ', '_')
        filename = f"{safe_title}_DRAFT_{timestamp}.pdf"
        output_path = str(Path(output_dir) / filename)
        
        return self.publish_to_pdf(
            story_id, 
            output_path, 
            include_synopsis=True,
            draft_watermark=True
        )
    
    def publish_final_draft(self, story_id: int, output_dir: str) -> str:
        """
        Publish a final draft PDF.
        Clean format without watermark.
        LOCKS the story from further editing.
        
        Returns:
            Path to the created PDF file
        """
        with DatabaseManager(self.db_path) as conn:
            story = Story.get_by_id(conn, story_id)
            if not story:
                raise ValueError(f"Story with ID {story_id} not found")
            
            # Update status to final published (locks it)
            story.publish_final(conn)
        
        # Generate filename
        safe_title = re.sub(r'[^\w\s-]', '', story.title).strip().replace(' ', '_')
        filename = f"{safe_title}_FINAL.pdf"
        output_path = str(Path(output_dir) / filename)
        
        return self.publish_to_pdf(
            story_id, 
            output_path, 
            include_synopsis=False,  # No summaries in final
            draft_watermark=False
        )
    
    def publish_to_docx(self, story_id: int, output_path: str,
                        include_synopsis: bool = False,
                        draft_watermark: bool = False) -> str:
        """
        Publish a story to DOCX format (Word document).
        Industry standard format for editors.
        
        Args:
            story_id: ID of the story to publish
            output_path: Path for the output DOCX file
            include_synopsis: Whether to include chapter summaries
            draft_watermark: Whether to add "DRAFT" header
            
        Returns:
            Path to the created DOCX file
        """
        with DatabaseManager(self.db_path) as conn:
            story = Story.get_by_id(conn, story_id)
            if not story:
                raise ValueError(f"Story with ID {story_id} not found")
            
            chapters = Chapter.get_by_story(conn, story_id)
        
        # Sort chapters by board position
        chapters.sort(key=lambda c: (c.board_y, c.board_x))
        
        # Create document
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(story.title)
        title_run.bold = True
        title_run.font.size = Pt(28)
        
        doc.add_paragraph()  # Spacer
        
        if story.synopsis:
            synopsis_para = doc.add_paragraph()
            synopsis_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            synopsis_run = synopsis_para.add_run(story.synopsis)
            synopsis_run.italic = True
            synopsis_run.font.size = Pt(11)
        
        if draft_watermark:
            doc.add_paragraph()
            draft_para = doc.add_paragraph()
            draft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            draft_run = draft_para.add_run(
                f"DRAFT - Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            draft_run.italic = True
        
        doc.add_page_break()
        
        # Chapters
        for i, chapter in enumerate(chapters):
            # Chapter heading
            chapter_num = i + 1
            heading = doc.add_heading(f"Chapter {chapter_num}: {chapter.title}", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Chapter summary
            if include_synopsis and chapter.summary:
                summary_para = doc.add_paragraph()
                summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                summary_run = summary_para.add_run(chapter.summary)
                summary_run.italic = True
                doc.add_paragraph()  # Spacer
            
            # Chapter content
            paragraphs = html_to_paragraphs(chapter.content)
            for j, para_text in enumerate(paragraphs):
                para = doc.add_paragraph(para_text)
                para.paragraph_format.first_line_indent = Inches(0.5) if j > 0 else Inches(0)
                para.paragraph_format.space_after = Pt(12)
            
            # Page break between chapters (except last)
            if i < len(chapters) - 1:
                doc.add_page_break()
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    def publish_to_epub(self, story_id: int, output_path: str,
                        author: str = "Unknown Author",
                        include_synopsis: bool = False) -> str:
        """
        Publish a story to EPUB format (ebook).
        Compatible with Amazon KDP, Apple Books, Kobo, etc.
        
        Args:
            story_id: ID of the story to publish
            output_path: Path for the output EPUB file
            author: Author name for metadata
            include_synopsis: Whether to include chapter summaries
            
        Returns:
            Path to the created EPUB file
        """
        with DatabaseManager(self.db_path) as conn:
            story = Story.get_by_id(conn, story_id)
            if not story:
                raise ValueError(f"Story with ID {story_id} not found")
            
            chapters = Chapter.get_by_story(conn, story_id)
        
        # Sort chapters by board position
        chapters.sort(key=lambda c: (c.board_y, c.board_x))
        
        # Create EPUB book
        book = epub.EpubBook()
        
        # Set metadata
        safe_title = re.sub(r'[^\w\s-]', '', story.title).strip()
        book.set_identifier(f'bluewriter-{story_id}-{datetime.now().strftime("%Y%m%d")}')
        book.set_title(story.title)
        book.set_language('en')
        book.add_author(author)
        
        # CSS for styling
        style_css = '''
        body { font-family: serif; line-height: 1.6; }
        h1 { text-align: center; margin-top: 2em; }
        h2 { text-align: center; }
        .chapter-title { text-align: center; font-size: 1.5em; margin: 2em 0 1em 0; }
        .synopsis { text-align: center; font-style: italic; margin-bottom: 2em; }
        p { text-indent: 1.5em; margin: 0.5em 0; }
        p.first { text-indent: 0; }
        .title-page { text-align: center; margin-top: 30%; }
        .title-page h1 { font-size: 2em; }
        .title-page .author { margin-top: 2em; }
        '''
        
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style_css
        )
        book.add_item(nav_css)
        
        # Title page
        title_content = f'''
        <html>
        <head><link rel="stylesheet" href="style/nav.css" /></head>
        <body>
        <div class="title-page">
            <h1>{story.title}</h1>
            <p class="author">by {author}</p>
            {f'<p class="synopsis">{story.synopsis}</p>' if story.synopsis else ''}
        </div>
        </body>
        </html>
        '''
        
        title_page = epub.EpubHtml(title='Title', file_name='title.xhtml', lang='en')
        title_page.content = title_content
        title_page.add_item(nav_css)
        book.add_item(title_page)
        
        # Create chapters
        epub_chapters = [title_page]
        spine = ['nav', title_page]
        
        for i, chapter in enumerate(chapters):
            chapter_num = i + 1
            
            # Build chapter HTML
            synopsis_html = ""
            if include_synopsis and chapter.summary:
                synopsis_html = f'<p class="synopsis">{chapter.summary}</p>'
            
            # Convert content to paragraphs
            paragraphs = html_to_paragraphs(chapter.content)
            content_html = ""
            for j, para_text in enumerate(paragraphs):
                css_class = 'first' if j == 0 else ''
                content_html += f'<p class="{css_class}">{para_text}</p>\n'
            
            chapter_html = f'''
            <html>
            <head><link rel="stylesheet" href="style/nav.css" /></head>
            <body>
            <h1 class="chapter-title">Chapter {chapter_num}: {chapter.title}</h1>
            {synopsis_html}
            {content_html}
            </body>
            </html>
            '''
            
            epub_chapter = epub.EpubHtml(
                title=f'Chapter {chapter_num}: {chapter.title}',
                file_name=f'chapter_{chapter_num:02d}.xhtml',
                lang='en'
            )
            epub_chapter.content = chapter_html
            epub_chapter.add_item(nav_css)
            
            book.add_item(epub_chapter)
            epub_chapters.append(epub_chapter)
            spine.append(epub_chapter)
        
        # Table of contents
        book.toc = [(epub.Section('Chapters'), epub_chapters[1:])]  # Skip title page
        
        # Add navigation files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Set spine
        book.spine = spine
        
        # Write EPUB file
        epub.write_epub(output_path, book, {})
        
        return output_path


def get_story_word_count(db_path: str, story_id: int) -> int:
    """Get total word count for a story."""
    with DatabaseManager(db_path) as conn:
        chapters = Chapter.get_by_story(conn, story_id)
    
    total_words = 0
    for chapter in chapters:
        text = html_to_text(chapter.content)
        total_words += len(text.split())
    
    return total_words


def get_story_stats(db_path: str, story_id: int) -> dict:
    """Get statistics for a story."""
    with DatabaseManager(db_path) as conn:
        story = Story.get_by_id(conn, story_id)
        chapters = Chapter.get_by_story(conn, story_id)
    
    total_words = 0
    total_chars = 0
    chapter_stats = []
    
    for chapter in chapters:
        text = html_to_text(chapter.content)
        words = len(text.split())
        chars = len(text)
        total_words += words
        total_chars += chars
        chapter_stats.append({
            'title': chapter.title,
            'words': words,
            'chars': chars
        })
    
    return {
        'title': story.title,
        'status': story.status,
        'chapter_count': len(chapters),
        'total_words': total_words,
        'total_chars': total_chars,
        'chapters': chapter_stats,
        'avg_chapter_words': total_words // len(chapters) if chapters else 0
    }
